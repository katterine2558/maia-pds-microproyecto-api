"""Genera el reporte de la Entrega 1 en formato .docx.

    python -m src.reportes.entrega1

El documento sale a docs/entregas/. Las cifras que aparecen en el texto estan
verificadas contra data/raw/diabetic_data.csv; las figuras las produce
src/data/exploracion.py y deben regenerarse antes que este documento.

"""

from __future__ import annotations

import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[2]
FIGURAS = RAIZ / "docs" / "entregas" / "figuras"
SALIDA = RAIZ / "docs" / "entregas" / "Entrega-1-reporte.docx"

REPO = "https://github.com/katterine2558/microproyecto-desarrollo-soluciones"
DATASET = "https://archive.ics.uci.edu/dataset/296/diabetes+130-us+hospitals+for+years+1999-2008"
REMOTO_DVC = "s3://maia-pds-diabetes-dvc-982005835034"

# Todo soporte se cita con su URL. Los archivos del repositorio se enlazan a main,
# que es la rama donde queda el estado entregado.
BLOB = f"{REPO}/blob/main"
TREE = f"{REPO}/tree/main"

TINTA = RGBColor(0x1A, 0x1A, 0x1A)
TINTA_2 = RGBColor(0x59, 0x59, 0x59)
ACENTO = RGBColor(0x14, 0x53, 0x8B)
PENDIENTE_COLOR = RGBColor(0x99, 0x33, 0x00)


# ---------------------------------------------------------------- utilidades

def git(*args: str) -> str:
    return subprocess.run(["git", *args], cwd=RAIZ, capture_output=True, text=True).stdout.strip()


def estado_git() -> dict:
    autores = [l.strip() for l in git("shortlog", "-sne", "--all").splitlines() if l.strip()]
    return {
        "commits": len(git("log", "--oneline", "--all").splitlines()),
        "autores": len(autores),
        "detalle_autores": autores,
        "ramas": [b.strip("* ").strip() for b in git("branch").splitlines()],
    }


def estilo_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TINTA
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for nombre, tam in (("Heading 1", 14), ("Heading 2", 11.5)):
        e = doc.styles[nombre]
        e.font.name = "Calibri"
        e.font.size = Pt(tam)
        e.font.bold = True
        e.font.color.rgb = TINTA
        e.paragraph_format.space_before = Pt(14 if tam > 12 else 10)
        e.paragraph_format.space_after = Pt(4)
        e.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def p(doc, texto: str = "", *, estilo: str | None = None, size: float | None = None,
      color: RGBColor | None = None, cursiva: bool = False, negrita: bool = False,
      espacio_antes: float | None = None, alineacion=None):
    """Parrafo con **negritas** inline usando marcadores dobles."""
    par = doc.add_paragraph(style=estilo) if estilo else doc.add_paragraph()
    if espacio_antes is not None:
        par.paragraph_format.space_before = Pt(espacio_antes)
    if alineacion is not None:
        par.paragraph_format.alignment = alineacion
    for i, trozo in enumerate(texto.split("**")):
        if not trozo:
            continue
        run = par.add_run(trozo)
        run.bold = negrita or (i % 2 == 1)
        run.italic = cursiva
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return par


def vineta(doc, texto: str):
    return p(doc, texto, estilo="List Bullet")


def pendiente(doc, texto: str):
    par = p(doc, texto, size=10, color=PENDIENTE_COLOR, cursiva=True)
    par.paragraph_format.left_indent = Cm(0.4)
    return par


def figura(doc, archivo: str, ancho_cm: float, leyenda: str):
    doc.add_picture(str(FIGURAS / archivo), width=Cm(ancho_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = p(doc, leyenda, size=8.5, color=TINTA_2, alineacion=WD_ALIGN_PARAGRAPH.CENTER)
    cap.paragraph_format.space_after = Pt(10)


def tabla(doc, encabezados: list[str], filas: list[list[str]], anchos: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for celda, titulo in zip(t.rows[0].cells, encabezados):
        celda.text = ""
        run = celda.paragraphs[0].add_run(titulo)
        run.bold = True
        run.font.size = Pt(9)
        celda.paragraphs[0].paragraph_format.space_after = Pt(2)
        _sombrear(celda, "EFEFEC")
    for fila in filas:
        celdas = t.add_row().cells
        for celda, valor in zip(celdas, fila):
            celda.text = ""
            par = celda.paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            for i, trozo in enumerate(str(valor).split("**")):
                if trozo:
                    run = par.add_run(trozo)
                    run.bold = i % 2 == 1
                    run.font.size = Pt(9)
    if anchos:
        for fila in t.rows:
            for celda, ancho in zip(fila.cells, anchos):
                celda.width = Cm(ancho)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def _sombrear(celda, hexcolor: str) -> None:
    tc = celda._tc.get_or_add_tcPr()
    sombra = OxmlElement("w:shd")
    sombra.set(qn("w:fill"), hexcolor)
    tc.append(sombra)


def codigo(doc, lineas: list[str]):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.5)
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, linea in enumerate(lineas):
        run = par.add_run(("\n" if i else "") + linea)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = TINTA_2
    return par


# ------------------------------------------------------------------ contenido

def construir() -> Document:
    g = estado_git()
    doc = Document()
    estilo_base(doc)

    s = doc.sections[0]
    s.page_height, s.page_width = Cm(29.7), Cm(21.0)
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.2)

    # ---- encabezado
    tit = p(doc, "Predicción de reingreso hospitalario temprano en pacientes diabéticos",
            negrita=True, size=16, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
    tit.paragraph_format.space_after = Pt(2)
    p(doc, "Entrega 1 · Micro-proyecto · Desarrollo de Soluciones · MAIA, Universidad de los Andes",
      size=10, color=TINTA_2, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, f"Repositorio: {REPO}", size=9.5, color=ACENTO, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
    p(doc, "Camilo Andrés Rodríguez Dueñas · Jasbyn Rainier Solano Carrillo · "
           "Leonardo Almanza Sánchez · Gineth Katerine Arias Carrillo",
      size=10, color=TINTA_2, alineacion=WD_ALIGN_PARAGRAPH.LEFT)

    # ---- 1
    doc.add_heading("1. Problema y su contexto", level=1)
    p(doc, "El reingreso hospitalario temprano —entendido en este proyecto como una nueva "
           "hospitalización dentro de los 30 días posteriores al alta— constituye el evento de "
           "interés porque representa una utilización repetida de servicios en un periodo corto "
           "y plantea una oportunidad de fortalecer el seguimiento posterior al egreso. En "
           "pacientes con diabetes, la necesidad de continuidad del cuidado es especialmente "
           "relevante por el carácter crónico de la enfermedad y por la coexistencia frecuente "
           "de múltiples diagnósticos, medicamentos y antecedentes de utilización hospitalaria. "
           "El propósito del proyecto no es atribuir una causa clínica al reingreso, sino "
           "identificar qué pacientes presentan mayor riesgo de volver al hospital en ese plazo.")

    p(doc, "Desde una perspectiva operativa, la utilidad del modelo está en apoyar una decisión "
           "de priorización antes de que el paciente abandone el hospital. Si la institución "
           "dispone de una capacidad limitada para realizar llamadas, controles u otras acciones "
           "de seguimiento, una estimación comparable de riesgo permite ordenar los egresos y "
           "concentrar primero esos recursos en los pacientes con mayor probabilidad estimada de "
           "reingreso.")

    p(doc, "Sin una estimación sistemática de riesgo, la priorización puede depender principalmente "
           "del criterio clínico individual y de las condiciones operativas del proceso de alta. "
           "En ese escenario, la capacidad de seguimiento no necesariamente se asigna primero a "
           "los pacientes con mayor riesgo estimado. Ese es el vacío que el prototipo busca "
           "abordar: no reemplazar el criterio clínico, sino **ordenar la lista** para apoyar la "
           "asignación de la capacidad disponible hacia los pacientes que podrían requerir mayor "
           "seguimiento.")

    p(doc, "El modelo propuesto es útil porque transforma información disponible durante la "
           "hospitalización —antecedentes de utilización de servicios, características del "
           "episodio, diagnósticos, pruebas y tratamiento registrado— en una estimación de riesgo "
           "individual. Esa estimación complementa, pero no sustituye, el criterio clínico: el "
           "prototipo ordena y prioriza casos para seguimiento; no autoriza altas, no diagnostica "
           "la causa del reingreso y no prescribe tratamiento.")

    # ---- 2
    doc.add_heading("2. Pregunta de negocio y alcance", level=1)
    pr = p(doc, "¿Qué pacientes diabéticos van a reingresar al hospital dentro de los 30 días "
                "siguientes al alta?", negrita=True)
    pr.paragraph_format.left_indent = Cm(0.4)
    p(doc, "La pregunta se responde el día del alta, antes de que el paciente salga, y su respuesta "
           "alimenta una decisión concreta: a qué pacientes se les programa control tras el egreso. "
           "El usuario es el personal de enfermería de la unidad de gestión hospitalaria.")
    p(doc, "El prototipo contempla dos formas complementarias de uso. La primera trabaja sobre el "
           "listado completo de egresos del día, ordenando a los pacientes según su probabilidad "
           "estimada de reingreso y permitiendo priorizar el seguimiento de acuerdo con la capacidad "
           "disponible. La segunda permite consultar individualmente a un paciente. Durante el diseño "
           "se descartaron dos alternativas más limitadas: una basada únicamente en consultas "
           "individuales, porque no permitiría establecer a quién priorizar cuando la capacidad de "
           "seguimiento es limitada; y otra basada únicamente en el listado diario, porque no "
           "permitiría evaluar un caso particular que requiera una consulta independiente.")       
    p(doc, "**Queda explícitamente por fuera del alcance.** El prototipo no estima la causa del "
           "reingreso, no sugiere tratamientos y no reemplaza el criterio clínico. Tampoco se conecta "
           "directamente a un sistema de información hospitalario: el listado de egresos se carga "
           "como archivo. La variable race se reserva para evaluar posteriormente el desempeño del "
           "modelo entre grupos y no se utiliza como predictora. De igual forma, payer_code se "
           "excluye de las variables predictoras, de acuerdo con las decisiones de preparación "
           "definidas a partir de la exploración de los datos.")
    p(doc, "**Límite de los datos.** Los registros corresponden a hospitales de Estados Unidos y "
           "cubren el periodo 1999–2008. Por esta razón, los patrones identificados no deben "
           "generalizarse automáticamente a una población hospitalaria actual o a una institución "
           "colombiana. El resultado de esta entrega debe entenderse como un prototipo metodológico; "
           "una implementación en otro contexto requeriría validar su desempeño con datos "
           "representativos de la población en la que se pretenda utilizar.")

    return doc, g


def parte_datos(doc, g):
    doc.add_heading("3. Conjuntos de datos a emplear", level=1)
    p(doc, f"Se emplea **Diabetes 130-US Hospitals for Years 1999-2008**, del UCI Machine Learning "
           f"Repository (dataset 296), publicado bajo licencia CC BY 4.0 y asociado al artículo de "
           f"Strack et al. (2014). Está disponible en {DATASET}.")
    p(doc, "Uno de los criterios considerados para seleccionar el conjunto de datos fue su "
           "disponibilidad inmediata. Este conjunto puede descargarse directamente desde UCI y cuenta "
           "con documentación pública, lo que permite trabajar con la información desde el inicio del "
           "proyecto. Frente a una alternativa que requiriera gestionar acceso institucional o "
           "recolectar información primaria, su disponibilidad resulta adecuada para el plazo "
           "establecido para el microproyecto (8 semanas).")
    p(doc, "Para que la procedencia sea verificable por un tercero, el repositorio incluye las "
           "huellas SHA-256 de los archivos junto con la URL exacta de descarga. Se contrastaron "
           "contra el archivo comprimido que UCI publica actualmente y coinciden byte a byte.")

    tabla(doc,
          ["Característica", "Valor"],
          [["Unidad de análisis", "Un encuentro hospitalario (no un paciente)"],
          ["Volumen", "101.766 encuentros × 50 columnas (47 predictoras, 2 identificadores y 1 objetivo)"],
           ["Pacientes únicos", "71.518 — hay pacientes con varios encuentros"],
           ["Periodo", "1999–2008, 130 hospitales de Estados Unidos"],
           ["Variable objetivo", "readmitted, con valores <30, >30 y NO"],
           ["Licencia", "CC BY 4.0"]],
          anchos=[5.0, 11.6])

    p(doc, "**Definición de la variable objetivo.** El campo readmitted contiene tres valores: "
           "<30, >30 y NO. Para responder la pregunta de negocio se construye una variable binaria: "
           "positivo cuando readmitted = <30 y negativo cuando readmitted = >30 o NO. Esta decisión "
           "no implica que los reingresos posteriores a 30 días carezcan de importancia clínica; "
           "simplemente quedan fuera de la ventana temporal que el proyecto busca predecir y "
           "priorizar.")
    p(doc, "**Familias de variables.** Las variables disponibles se agrupan en demográficas (edad, sexo, "
           "etnia), administrativas (tipo de admisión, destino al egreso, aseguradora, especialidad "
           "que da el alta), de utilización previa (consultas ambulatorias, urgencias y "
           "hospitalizaciones del año anterior), clínicas (días de estancia, número de "
           "diagnósticos, códigos ICD-9, resultados de laboratorio) y farmacológicas (23 columnas, "
           "una por medicamento). El repositorio incluye el diccionario completo de las 50 "
           "variables, con su tipo y sus valores.")


def parte_repos(doc, g):
    doc.add_heading("4. Repositorio Git en uso para el código", level=1)
    p(doc, f"El código está en {REPO}. El repositorio sigue un flujo git-flow adaptado al proyecto: "
           f"main conserva únicamente los estados entregados y se etiqueta en cada entrega; develop "
           f"integra el trabajo diario; y cada ítem de trabajo vive en su propia rama feature/*, "
           f"que sale de develop y vuelve a develop mediante merges que preservan la autoría —no se "
           f"usa squash ni rebase, precisamente porque los commits son la evidencia de la "
           f"contribución individual.")
    codigo(doc, [
        "$ git shortlog -sne --all",
        *(f"  {a}" for a in g["detalle_autores"]),
        "",
        f"$ git log --oneline --all | wc -l      →  {g['commits']} commits",
        f"$ git branch                           →  {', '.join(g['ramas'])}",
    ])
    
    p(doc, "El repositorio es público, de modo que el equipo de tutores puede revisar el código, "
           "el historial y las ramas sin solicitar acceso. Los cuatro integrantes figuran como "
           "colaboradores con permiso de escritura e integran su trabajo mediante merges directos "
           "a develop; cualquier persona ajena al equipo solo puede proponer cambios vía fork y "
           "pull request, que es el comportamiento por defecto de GitHub y no requiere reglas de "
           "protección de rama.")

    doc.add_heading("5. Repositorio DVC en uso para los datos", level=1)
    p(doc, f"Los datos se versionan con DVC sobre el mismo repositorio de Git. En Git viaja "
           f"únicamente el puntero data/raw.dvc —cinco líneas de texto— mientras que los archivos "
           f"reales, 18 MB, viven en el remoto {REMOTO_DVC}. La separación es explícita en el "
           f".gitignore: el contenido de data/ está excluido y solo se permiten los punteros .dvc.")
    p(doc, "El remoto tiene lectura pública, de modo que cualquier integrante clona el repositorio "
           "y ejecuta dvc pull sin necesidad de credenciales. Se requieren credenciales solo para "
           "escribir, y se configuran localmente en un archivo que nunca se versiona.")
    codigo(doc, [
        "$ cat data/raw.dvc",
        "  outs:",
        "  - md5: c96228354d9545119486617492cc9c37.dir",
        "    size: 19161930",
        "    nfiles: 2",
        "    path: raw",
        "",
        "$ dvc status -c",
        "  Cache and remote 'storage' are in sync.",
    ])
    p(doc, "El puntero versionado en Git junto al CSV ausente del repositorio prueba las dos cosas "
           "a la vez: los datos están versionados por DVC y están fuera de Git.", size=9.5,
      color=TINTA_2)


def parte_exploracion(doc):
    doc.add_heading("6. Exploración de los datos", level=1)

    doc.add_heading("6.1 Los faltantes engañan en las dos direcciones", level=2)
    p(doc, "El primer hallazgo no es sobre los pacientes sino sobre el archivo, y condiciona todo "
           "lo demás. El centinela de dato faltante en este conjunto es el carácter ?, no un valor "
           "nulo. Con las opciones por defecto de pandas, la instrucción habitual para contar "
           "faltantes devuelve cero en todas las columnas, y el conjunto parece completo. No lo "
           "está: weight está vacía en el 96,9 % de las filas.")
    p(doc, "El error simétrico es más peligroso porque es silencioso. Las columnas A1Cresult y "
           "max_glu_serum traen el texto literal None, que pandas incluye en su lista de valores "
           "nulos por defecto y convierte sin avisar. Pero ahí None no es un faltante: significa "
           "que **no se ordenó el examen**, y eso es información clínica. Que a un paciente "
           "diabético no le midan la hemoglobina glicosilada durante la hospitalización dice algo "
           "sobre cómo se manejó su caso, y es justamente la hipótesis del artículo original. "
           "Tratarlo como dato ausente e imputarlo destruiría 84.748 observaciones válidas.")
    tabla(doc,
          ["Columna", "Lectura por defecto", "Realidad"],
          [["race", "0 nulos", "2.273 valores ? (2,2 %)"],
           ["weight", "0 nulos", "98.569 valores ? (96,9 %)"],
           ["medical_specialty", "0 nulos", "49,1 % ausente"],
           ["A1Cresult", "84.748 nulos", "0 faltantes — None = examen no ordenado"],
           ["max_glu_serum", "96.420 nulos", "0 faltantes — None = examen no ordenado"]],
          anchos=[4.2, 4.6, 7.8])
    p(doc, "**Por lo tanto**, todo el pipeline lee el archivo desactivando la interpretación "
           "automática de nulos y convierte ? explícitamente, columna por columna. La ausencia de "
           "especialidad, con casi la mitad de los registros, se tratará como categoría propia y no "
           "como dato a imputar.")

    doc.add_heading("6.2 Exclusiones: dos motivos distintos", level=2)
    p(doc, "El destino al egreso identifica encuentros que no deben entrar al análisis, por razones "
           "que conviene no confundir. Los 1.652 encuentros de pacientes fallecidos no pueden "
           "reingresar, y los datos lo confirman: ninguno registra reingreso. Excluirlos corrige un "
           "imposible. Los 771 con egreso a hospicio sí reingresan —43 lo hacen antes de los 30 "
           "días— y se excluyen por "
           "una razón distinta: son pacientes en cuidado de fin de vida, donde agendar un control "
           "para evitar el reingreso no es la intervención que el tablero decide.")
    p(doc, "En conjunto salen 2.423 encuentros, el 2,38 %. La base de trabajo queda en **99.343 "
           "egresos**, con una tasa de reingreso temprano del **11,4 %**.")
    p(doc, "**Por lo tanto**, con menos de doce reingresos tempranos por cada cien casos, la exactitud utilizada de "
           "manera aislada resulta insuficiente para evaluar el desempeño del modelo. Por ejemplo, "
           "un modelo que predijera siempre «no reingresa» alcanzaría una exactitud cercana al "
           "88,6 %, aun sin identificar ningún reingreso temprano. Por esta razón, la evaluación "
           "deberá considerar métricas que permitan valorar la identificación y priorización de "
           "la clase minoritaria.")

    doc.add_heading("6.3 Dónde se concentra el riesgo", level=2)
    p(doc, "La utilización previa es, con diferencia, la señal más fuerte del conjunto. La tasa de "
           "reingreso pasa de 8,6 % en quienes no estuvieron hospitalizados el año anterior a "
           "37,1 % en quienes lo estuvieron cinco veces o más: un gradiente de más de cuatro veces, "
           "monótono en todos los tramos.")
    figura(doc, "01-reingreso-por-hospitalizaciones-previas.png", 16.0,
           "Figura 1. Tasa de reingreso < 30 días según hospitalizaciones en el año previo. "
           "Base: 99.343 egresos tras exclusiones.")
    p(doc, "El patrón observado convierte a number_inpatient y a las demás variables de utilización "
           "previa en candidatas relevantes para el modelamiento. Su aporte definitivo deberá "
           "confirmarse durante el entrenamiento y la validación del modelo. La maqueta las muestra "
           "junto a la probabilidad estimada porque son antecedentes fáciles de interpretar para "
           "el usuario.")

    figura(doc, "02-reingreso-por-especialidad.png", 15.0,
           "Figura 2. Tasa de reingreso por especialidad que da el alta. Solo especialidades con al "
           "menos 500 egresos; se excluye la categoría sin especialidad registrada.")
    p(doc, "Entre especialidades el rango va de 4,8 % en ginecobstetricia a 16,1 % en nefrología, "
           "que dobla a cardiología (8,0 %). **Por lo tanto**, hay servicios donde reforzar la "
           "continuidad del cuidado rinde más, y esa es una decisión que no necesita predecir a "
           "ningún paciente en particular: alimenta el componente descriptivo del tablero.")

    figura(doc, "03-reingreso-por-edad.png", 16.0,
           "Figura 3. Tasa de reingreso por grupo de edad. El pico de 20–30 años se apoya en 1.649 "
           "registros, frente a más de 20.000 en los grupos mayores.")
    p(doc, "La edad muestra un ascenso sostenido a partir de los 50 años, pero el máximo aparente "
           "está en el grupo de 20 a 30. **Por lo tanto**, conviene leerlo con reservas: descansa "
           "sobre 1.649 registros, un orden de magnitud menos que los grupos mayores, y no "
           "justificaría por sí solo una regla de priorización.")

    doc.add_heading("6.4 Riesgo de fuga por pacientes repetidos", level=2)
    p(doc, "Los 99.343 encuentros corresponden a 69.990 pacientes distintos: 16.341 pacientes "
           "aparecen más de una vez. **Por lo tanto**, una partición aleatoria en entrenamiento y "
           "prueba dejaría encuentros del mismo paciente a ambos lados y produciría un desempeño "
           "optimista que no se sostendría en operación. La partición se hará por patient_nbr.")
    p(doc, "Se detectaron además dos columnas de varianza cero —examide y citoglipton, con un único "
           "valor en las 101.766 filas— y otras siete de medicamentos con menos de cincuenta "
           "registros distintos de «no recetado». Todas se descartan.")


def parte_maqueta(doc):
    doc.add_heading("7. Maqueta del prototipo", level=1)
    p(doc, "La maqueta define tres pantallas y trece elementos identificados de E1 a E13. Las cifras "
           "que muestra el componente descriptivo son reales, calculadas sobre los 99.343 egresos; "
           "los datos de las pantallas de predicción son ilustrativos.")
    figura(doc, "04-maqueta-pantalla-1-priorizacion.png", 16.6,
           "Figura 4. Pantalla 1 — priorización de los egresos del día. La enfermera declara su "
           "capacidad (E1) y la línea de corte se mueve con ese número (E3).")
    p(doc, "La pantalla de priorización es la que sostiene la decisión. Ordena por probabilidad "
           "estimada en vez de emitir una etiqueta, porque con una tasa base de 11,4 % lo que sirve "
           "a quien programa los controles es el orden, no un sí o un no. La línea de capacidad "
           "vuelve visible una limitación que hoy nadie mide: cuántos pacientes de riesgo alto "
           "quedan sin control por falta de recurso.")
    figura(doc, "07-arquitectura.png", 15.5,
           "Figura 5. Arquitectura. El tablero consume las predicciones por HTTP; nunca carga el "
           "artefacto del modelo.")
    p(doc, "Los tres componentes se despliegan en contenedores separados. El tablero obtiene las "
           "predicciones llamando a la API, nunca abriendo el archivo del modelo: si lo hiciera, "
           "los dos quedarían amarrados y no sería posible actualizar el modelo sin volver a "
           "desplegar el tablero.")
    p(doc, "**Relación con la pregunta de negocio.** Cada elemento de la maqueta responde una parte "
           "concreta de la pregunta; la memoria completa incluye la tabla de trazabilidad de los "
           "trece. Un extracto:")
    tabla(doc,
          ["Id", "Elemento", "Qué responde", "Tipo"],
          [["E2", "Orden por probabilidad estimada", "¿A quién atiendo primero?", "Predictivo"],
           ["E3", "Línea de capacidad visible", "¿A cuántos alcanzo hoy?", "Predictivo"],
           ["E7", "Probabilidad calibrada", "¿Qué tan probable es este caso?", "Predictivo"],
           ["E9", "Factores que pesaron", "¿Por qué este paciente?", "Predictivo"],
           ["E10", "Reingreso por hospitalizaciones previas", "¿Dónde está el riesgo?", "Descriptivo"],
           ["E11", "Reingreso por servicio", "¿Dónde reforzar el cuidado?", "Descriptivo"]],
          anchos=[1.2, 6.0, 6.2, 3.2])
    p(doc, "La maqueta se iteró durante la semana 3 con los datos ya explorados. El cambio principal "
           "fue anclar las bandas de riesgo en la tasa general observada —bajo por debajo de 11,4 %, "
           "medio hasta el doble, alto de 22,8 % en adelante— en lugar de usar percentiles, que se "
           "moverían con la composición del listado de cada día.")


def parte_equipo(doc):
    doc.add_heading("8. Reporte de trabajo en equipo", level=1)
    
    tabla(doc,
          ["Actividad", "Responsable"],
          [["Problema y contexto", "Todos"],
           ["Pregunta de negocio y alcance", "Todos"],
           ["Descripción del conjunto de datos", "Gineth Katerine Arias Carrillo"],
           ["Exploración de los datos", "Gineth Katerine Arias Carrillo"],
           ["Maqueta del prototipo", "Gineth Katerine Arias Carrillo, Jasbyn Rainier Solano Carrillo"],
           ["Repositorios creados", "Gineth Katerine Arias Carrillo"],
           ["Revisión y ajuste del reporte", "Camilo Rodríguez Dueñas"]],
          anchos=[10.5, 6.1])

    doc.add_heading("9. Referencias", level=1)

    p(doc, "Clore, J., Cios, K., DeShazo, J., & Strack, B. (2014). "
           "Diabetes 130-US Hospitals for Years 1999-2008 [Dataset]. "
           "UCI Machine Learning Repository. https://doi.org/10.24432/C5230J",
      size=9.5)

    p(doc, "Strack, B., DeShazo, J. P., Gennings, C., Olmo, J. L., Ventura, S., "
           "Cios, K. J., & Clore, J. N. (2014). Impact of HbA1c measurement on "
           "hospital readmission rates: Analysis of 70,000 clinical database "
           "patient records. BioMed Research International, 2014, Article 781670. "
           "https://doi.org/10.1155/2014/781670",
      size=9.5)
    
    doc.add_heading("Soportes de la entrega", level=1)
    for s in [
        f"Repositorio, con el historial de commits y las ramas: {REPO}",
        f"Conjunto de datos original, UCI 296, licencia CC BY 4.0: {DATASET}",
        f"Remoto DVC de los datos, lectura pública: {REMOTO_DVC}",
        f"Notebook de exploración, nueve secciones ejecutadas: {BLOB}/notebooks/eda.ipynb",
        f"Diccionario de las 50 variables: {BLOB}/docs/diccionario-variables.md",
        f"Huellas SHA-256 y URL de descarga de los datos: "
        f"{BLOB}/docs/soportes/checksums-datos-crudos.txt",
        f"Maqueta, pantallas: {BLOB}/docs/maqueta/pantallas.html",
        f"Maqueta, memoria de los trece elementos: {BLOB}/docs/maqueta/memoria.html",
        f"Figuras en resolución completa: {TREE}/docs/entregas/figuras",
        f"Política IAM del remoto de datos: {BLOB}/docs/soportes/politica-iam-dvc.json",
    ]:
        vineta(doc, s)


def main() -> int:
    doc, g = construir()
    parte_datos(doc, g)
    parte_repos(doc, g)
    parte_exploracion(doc)
    parte_maqueta(doc)
    parte_equipo(doc)
    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SALIDA)
    print(f"{SALIDA.relative_to(RAIZ)}")
    print(f"commits detectados: {g['commits']} | autores: {g['autores']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
